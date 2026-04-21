#!/usr/bin/env python3
"""
Convert Markdown file to DOCX with proper formatting
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def parse_markdown_to_docx(md_file, docx_file):
    """Parse markdown file and create formatted DOCX"""

    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Split into lines
    lines = content.split('\n')

    # Track current list level
    in_list = False
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    for i, line in enumerate(lines):
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_lines:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Handle table rows
        if line.strip().startswith('|') and '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            # Check if next line is not a table row
            if i + 1 >= len(lines) or not lines[i + 1].strip().startswith('|'):
                # Create table
                create_table(doc, table_lines)
                in_table = False
                table_lines = []
            continue

        # Skip horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph()
            continue

        # Handle headers
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()

            # Remove emoji if present (optional)
            text = text.strip()

            p = doc.add_paragraph(text)

            if level == 1:
                p.style = 'Heading 1'
                for run in p.runs:
                    run.font.size = Pt(24)
                    run.font.color.rgb = RGBColor(204, 0, 0)  # Red
                    run.bold = True
            elif level == 2:
                p.style = 'Heading 2'
                for run in p.runs:
                    run.font.size = Pt(18)
                    run.font.color.rgb = RGBColor(204, 0, 0)
                    run.bold = True
            elif level == 3:
                p.style = 'Heading 3'
                for run in p.runs:
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.bold = True
            elif level == 4:
                p.style = 'Heading 4'
                for run in p.runs:
                    run.font.size = Pt(12)
                    run.bold = True
            continue

        # Handle bullet lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            text = format_inline_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
            continue

        # Handle numbered lists
        if re.match(r'^\s*\d+\.\s', line):
            text = re.sub(r'^\s*\d+\.\s', '', line)
            text = format_inline_markdown(text)
            p = doc.add_paragraph(text, style='List Number')
            continue

        # Handle bold/italic
        if line.strip():
            text = format_inline_markdown(line.strip())
            p = doc.add_paragraph(text)
            p.style = 'Normal'
        else:
            # Empty line - add spacing
            doc.add_paragraph()

def create_table(doc, table_lines):
    """Create a formatted table from markdown table lines"""
    # Remove separator line (the one with dashes)
    rows = [line for line in table_lines if not re.match(r'^\s*\|[\s\-:]+\|', line)]

    if not rows:
        return

    # Parse table data
    parsed_rows = []
    for row in rows:
        cells = [cell.strip() for cell in row.split('|')[1:-1]]
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    # Create table
    num_cols = len(parsed_rows[0])
    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Populate table
    for i, row_data in enumerate(parsed_rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            # Format header row
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                # Try to set background color
                try:
                    from docx.oxml import OxmlElement
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'CC0000')
                    cell._element.get_or_add_tcPr().append(shading_elm)
                except:
                    pass  # Skip if can't set background

def format_inline_markdown(text):
    """Format inline markdown (bold, italic, code)"""
    # Handle code inline
    text = re.sub(r'`([^`]+)`', r'\1', text)

    # Handle bold
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'__([^_]+)__', r'\1', text)

    # Handle italic
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)

    return text

if __name__ == '__main__':
    input_file = '/Users/zhaight/Desktop/Value/FUNCTIONALITY_SUMMARY.md'
    output_file = '/Users/zhaight/Desktop/Value/FUNCTIONALITY_SUMMARY.docx'

    print(f"Converting {input_file} to {output_file}...")
    parse_markdown_to_docx(input_file, output_file)
    print(f"✅ Successfully created {output_file}")
