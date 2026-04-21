#!/usr/bin/env python3
"""
Simple Markdown to DOCX converter
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_md_to_docx(input_file, output_file):
    """Convert markdown to DOCX with basic formatting"""

    # Read markdown
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create document
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    in_code = False
    in_table = False
    table_data = []

    for line in lines:
        line = line.rstrip()

        # Code blocks
        if line.startswith('```'):
            in_code = not in_code
            continue

        if in_code:
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            continue

        # Tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            # Skip separator lines
            if re.match(r'^\|[\s\-:]+\|', line):
                continue
            table_data.append(line)
            continue
        elif in_table:
            # End of table
            add_table(doc, table_data)
            in_table = False
            table_data = []

        # Skip horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph()
            continue

        # Headers
        if line.startswith('#'):
            level = 0
            while line[level] == '#':
                level += 1
            text = line[level:].strip()

            p = doc.add_paragraph(text)
            if level == 1:
                p.style = 'Heading 1'
                for run in p.runs:
                    run.font.size = Pt(24)
                    run.font.color.rgb = RGBColor(204, 0, 0)
            elif level == 2:
                p.style = 'Heading 2'
                for run in p.runs:
                    run.font.size = Pt(18)
                    run.font.color.rgb = RGBColor(204, 0, 0)
            elif level == 3:
                p.style = 'Heading 3'
                for run in p.runs:
                    run.font.size = Pt(14)
            elif level == 4:
                p.style = 'Heading 4'
                for run in p.runs:
                    run.font.size = Pt(12)
            continue

        # Lists
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = clean_text(line.strip()[2:])
            doc.add_paragraph(text, style='List Bullet')
            continue

        if re.match(r'^\s*\d+\.\s', line):
            text = clean_text(re.sub(r'^\s*\d+\.\s', '', line))
            doc.add_paragraph(text, style='List Number')
            continue

        # Normal text
        if line.strip():
            text = clean_text(line.strip())
            doc.add_paragraph(text)
        else:
            doc.add_paragraph()

    # Handle remaining table
    if in_table and table_data:
        add_table(doc, table_data)

    # Save
    doc.save(output_file)
    print(f"✅ Successfully created {output_file}")

def add_table(doc, table_lines):
    """Add a table to the document"""
    if not table_lines:
        return

    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)

    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = clean_text(cell_text)
                # Bold header row
                if i == 0:
                    for para in table.rows[i].cells[j].paragraphs:
                        for run in para.runs:
                            run.bold = True

def clean_text(text):
    """Remove markdown formatting"""
    # Remove bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # Remove italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Remove inline code
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

if __name__ == '__main__':
    input_file = '/Users/zhaight/Desktop/Value/FUNCTIONALITY_SUMMARY.md'
    output_file = '/Users/zhaight/Desktop/Value/FUNCTIONALITY_SUMMARY.docx'

    print(f"Converting {input_file}...")
    convert_md_to_docx(input_file, output_file)
